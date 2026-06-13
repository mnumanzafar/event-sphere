"""
Verify all changes by extracting text from modified chapters and counting content
"""
from docx import Document
import os

report_dir = r"E:\FYP-main\report_extracted\report"
backup_dir = r"E:\FYP-main\report_extracted\report_backup"

print("=" * 70)
print("VERIFICATION: Comparing original vs modified chapter sizes")
print("=" * 70)

total_original = 0
total_modified = 0

for fname in sorted(os.listdir(report_dir)):
    if not fname.endswith('.docx'):
        continue
    
    # Modified version
    mod_doc = Document(os.path.join(report_dir, fname))
    mod_paras = [p.text.strip() for p in mod_doc.paragraphs if p.text.strip()]
    mod_chars = sum(len(t) for t in mod_paras)
    mod_tables = len(mod_doc.tables)
    
    # Original version
    orig_path = os.path.join(backup_dir, fname)
    if os.path.exists(orig_path):
        orig_doc = Document(orig_path)
        orig_paras = [p.text.strip() for p in orig_doc.paragraphs if p.text.strip()]
        orig_chars = sum(len(t) for t in orig_paras)
        orig_tables = len(orig_doc.tables)
        
        total_original += orig_chars
        total_modified += mod_chars
        
        delta_paras = len(mod_paras) - len(orig_paras)
        delta_chars = mod_chars - orig_chars
        
        print(f"\n{fname}:")
        print(f"  Original: {len(orig_paras)} paras, {orig_chars:,} chars, {orig_tables} tables")
        print(f"  Modified: {len(mod_paras)} paras, {mod_chars:,} chars, {mod_tables} tables")
        print(f"  Change:   {delta_paras:+d} paras, {delta_chars:+,d} chars")
    else:
        print(f"\n{fname}: No backup found for comparison")

print(f"\n{'=' * 70}")
print(f"TOTAL CHANGE: {total_modified - total_original:+,d} characters")
reduction_pct = ((total_original - total_modified) / total_original * 100) if total_original > 0 else 0
print(f"Text reduction: {reduction_pct:.1f}%")
# Rough estimate: ~2500 chars/page in a formatted report
est_original_pages = total_original / 2500
est_modified_pages = total_modified / 2500
print(f"Estimated pages: ~{est_original_pages:.0f} -> ~{est_modified_pages:.0f}")
print(f"{'=' * 70}")

# Show key content verification
print(f"\n{'=' * 70}")
print("CONTENT VERIFICATION: Checking key additions exist")
print(f"{'=' * 70}")

checks = {
    "ch5.docx": [
        "5.4 Environment Configuration",
        "5.5 State Management",
        "5.6 Expanded Role-Based",
        "5.7 Gamification Module",
        "5.8 AI Chatbot Assistant",
        "5.9 Social Sharing",
        "5.10 Additional Implemented",
        "5.11 Database Security",
        "flutter_dotenv",
        "Riverpod",
        "Super Admin",
        "rules-based intent",
        "WhatsApp",
        "Row Level Security",
    ],
    "ch4.docx": [
        "Vice President Interface",
        "President Interface", 
        "EventFeedback",
        "UserPoints",
        "EventWaitlist",
        "RoleRequests",
        "five different role-based",
    ],
    "EventSphere_Ch7.docx": [
        "Limited Offline Functionality",
        "Hive local storage",
        "five user roles",
    ],
}

for fname, expected_items in checks.items():
    fpath = os.path.join(report_dir, fname)
    doc = Document(fpath)
    full_text = ' '.join(p.text for p in doc.paragraphs)
    
    print(f"\n{fname}:")
    for item in expected_items:
        found = item in full_text
        status = "[OK]" if found else "[MISSING]"
        print(f"  {status} '{item}'")
