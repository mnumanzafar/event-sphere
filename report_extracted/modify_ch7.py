"""
Chapter 7 - CONCLUSIONS, LIMITATIONS AND FUTURE IMPROVEMENTS
TRIM: Convert verbose future improvements to concise format
FIX: Update limitation about "no offline caching" (project has it)
FIX: Update role references from 3 to 5
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document(r"E:\FYP-main\report_extracted\report\EventSphere_Ch7.docx")

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)

def make_paragraph(text, style='Normal'):
    new_p = OxmlElement('w:p')
    new_ppr = OxmlElement('w:pPr')
    new_style_el = OxmlElement('w:pStyle')
    new_style_el.set(qn('w:val'), style)
    new_ppr.append(new_style_el)
    new_p.append(new_ppr)
    new_run = OxmlElement('w:r')
    new_text = OxmlElement('w:t')
    new_text.set(qn('xml:space'), 'preserve')
    new_text.text = text
    new_run.append(new_text)
    new_p.append(new_run)
    return new_p

paragraphs = doc.paragraphs

# Map all sections
sections = {}
for i, p in enumerate(paragraphs):
    txt = p.text.strip()
    if txt.startswith('7.1'):
        sections['7.1_start'] = i
    elif txt.startswith('7.2 Limitations'):
        sections['7.2_start'] = i
    elif txt.startswith('7.2.1'):
        sections['7.2.1_start'] = i
    elif txt.startswith('7.2.2'):
        sections['7.2.2_start'] = i
    elif txt.startswith('7.2.3'):
        sections['7.2.3_start'] = i
    elif txt.startswith('7.2.4'):
        sections['7.2.4_start'] = i
    elif txt.startswith('7.3 Future'):
        sections['7.3_start'] = i
    elif txt.startswith('7.3.1'):
        sections['7.3.1_start'] = i
    elif txt.startswith('7.3.2'):
        sections['7.3.2_start'] = i
    elif txt.startswith('7.3.3'):
        sections['7.3.3_start'] = i
    elif txt.startswith('7.3.4'):
        sections['7.3.4_start'] = i
    elif txt.startswith('7.3.5'):
        sections['7.3.5_start'] = i
    elif txt.startswith('7.3.6'):
        sections['7.3.6_start'] = i
    elif txt.startswith('7.3.7'):
        sections['7.3.7_start'] = i
    elif txt.startswith('7.3.8'):
        sections['7.3.8_start'] = i

print(f"Found sections: {sections}")

# ============================================================================
# STEP 1: FIX limitation 7.2.4 - "No Offline Event Caching"
# The project DOES have offline caching via Hive/CacheService/OfflineService
# Change this to a more accurate limitation
# ============================================================================
if '7.2.4_start' in sections:
    heading_p = paragraphs[sections['7.2.4_start']]
    # Change heading text
    for run in heading_p.runs:
        if 'No Offline Event Caching' in run.text:
            run.text = run.text.replace('No Offline Event Caching', 'Limited Offline Functionality')
        elif 'Offline' in run.text and 'Caching' in run.text:
            run.text = '7.2.4 Limited Offline Functionality'
    
    # Change description paragraph (next paragraph after heading)
    desc_idx = sections['7.2.4_start'] + 1
    if desc_idx < len(paragraphs):
        desc_p = paragraphs[desc_idx]
        new_text = (
            "While Event Sphere includes a basic offline caching layer using Hive local storage "
            "(CacheService) and a connectivity-aware action queue (OfflineService), the offline "
            "functionality is limited to caching previously loaded event data. Users cannot create "
            "new events, register for events, or perform QR attendance scanning while offline. Full "
            "offline mode with complete read-write synchronization remains a future enhancement."
        )
        if desc_p.runs:
            desc_p.runs[0].text = new_text
            for r in desc_p.runs[1:]:
                r.text = ''
    
    print("Fixed limitation 7.2.4 - updated offline caching description")

# ============================================================================
# STEP 2: TRIM future improvements (7.3.1 - 7.3.8) into condensed format
# Replace each verbose subsection with 2-3 sentence descriptions
# ============================================================================
future_sections = [
    ('7.3.1_start', '7.3.2_start', 
     "7.3.1 iOS Platform Support",
     "Extending Event Sphere to iOS using Flutter's cross-platform capabilities. This involves "
     "configuring Xcode build settings, Firebase iOS SDK integration, Apple Push Notification "
     "Service setup, and App Store submission compliance with Apple's Human Interface Guidelines."),
    
    ('7.3.2_start', '7.3.3_start',
     "7.3.2 Advanced AI Integration",
     "Replacing the current rules-based chatbot with a Google Gemini AI-powered assistant capable "
     "of natural language understanding, personalized event recommendations based on attendance "
     "history, and intelligent event summarization for quick user consumption."),
    
    ('7.3.3_start', '7.3.4_start',
     "7.3.3 Multi-Language Support",
     "Implementing Flutter's localization framework (flutter_localizations, intl) to support Urdu, "
     "Arabic, and Chinese alongside English, with right-to-left layout support for Urdu and Arabic."),
    
    ('7.3.4_start', '7.3.5_start',
     "7.3.4 Payment Integration",
     "Integrating Stripe or JazzCash payment gateways for paid event ticketing, with automated "
     "receipt generation, refund processing, and expense report integration."),
    
    ('7.3.5_start', '7.3.6_start',
     "7.3.5 Advanced Analytics Dashboard",
     "Expanding the admin dashboard with predictive analytics using historical attendance patterns, "
     "real-time event popularity heatmaps, automated engagement reports, and exportable BI-style "
     "visualizations."),
    
    ('7.3.6_start', '7.3.7_start',
     "7.3.6 Calendar Integration",
     "Implementing native calendar synchronization to add registered events to Google Calendar or "
     "Apple Calendar with automated reminders and schedule conflict detection."),
    
    ('7.3.7_start', '7.3.8_start',
     "7.3.7 Web Application Version",
     "Deploying a Flutter Web companion application using Flutter's web rendering engine for "
     "desktop access, enabling broader reach for event browsing and administrative tasks."),
]

for start_key, end_key, new_heading, new_desc in future_sections:
    if start_key in sections and end_key in sections:
        # Delete all paragraphs between this section and next
        to_delete = []
        for i in range(sections[start_key], sections[end_key]):
            to_delete.append(i)
        
        # Insert condensed replacement
        ref_para = paragraphs[sections[start_key]]
        parent = ref_para._element.getparent()
        insert_pos = list(parent).index(ref_para._element)
        
        desc_p = make_paragraph(new_desc, 'Normal')
        head_p = make_paragraph(new_heading, 'Heading3')
        
        parent.insert(insert_pos, desc_p)
        parent.insert(insert_pos, head_p)
        
        # Delete old paragraphs
        for i in sorted(to_delete, reverse=True):
            try:
                delete_paragraph(paragraphs[i])
            except:
                pass

# Handle last section 7.3.8 separately
if '7.3.8_start' in sections:
    # Find end of chapter (last meaningful paragraph)
    last_meaningful = sections['7.3.8_start']
    for i in range(sections['7.3.8_start'], len(paragraphs)):
        if paragraphs[i].text.strip():
            last_meaningful = i
    
    to_delete_last = []
    for i in range(sections['7.3.8_start'], last_meaningful + 1):
        to_delete_last.append(i)
    
    ref_para = paragraphs[sections['7.3.8_start']]
    parent = ref_para._element.getparent()
    insert_pos = list(parent).index(ref_para._element)
    
    desc_p = make_paragraph(
        "Implementing an in-app virtual event hosting module with live streaming using WebRTC, "
        "interactive Q&A sessions, virtual breakout rooms, and real-time polls during live events "
        "to support hybrid physical-virtual event formats.",
        'Normal'
    )
    head_p = make_paragraph("7.3.8 Virtual Event Hosting", 'Heading3')
    
    parent.insert(insert_pos, desc_p)
    parent.insert(insert_pos, head_p)
    
    for i in sorted(to_delete_last, reverse=True):
        try:
            delete_paragraph(paragraphs[i])
        except:
            pass

print("Trimmed all 8 future improvement sections")

# ============================================================================
# STEP 3: Fix role references in conclusions (7.1) 
# ============================================================================
for p in doc.paragraphs:
    for run in p.runs:
        if 'three user roles' in run.text:
            run.text = run.text.replace('three user roles', 'five user roles')
        if 'three types of users' in run.text:
            run.text = run.text.replace('three types of users', 'five types of users')
        if 'three role-based' in run.text and 'three-layer' not in run.text:
            run.text = run.text.replace('three role-based', 'five role-based')
        if '(Admin, Faculty, Student)' in run.text:
            run.text = run.text.replace(
                '(Admin, Faculty, Student)',
                '(Student, Vice President, President, Admin, Super Admin)'
            )
        if 'Admin, Faculty, and Student' in run.text:
            run.text = run.text.replace(
                'Admin, Faculty, and Student',
                'Student, Vice President, President, Admin, and Super Admin'
            )

print("Fixed role references in Chapter 7")

doc.save(r"E:\FYP-main\report_extracted\report\EventSphere_Ch7.docx")
print("Chapter 7 saved successfully!")
