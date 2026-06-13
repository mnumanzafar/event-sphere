"""
Chapter 3 - REQUIREMENT ANALYSIS
TRIM: Consolidate 5 verbose use case description tables into brief summary,
      Remove repetitive NFR intro sentences
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document(r"E:\FYP-main\report_extracted\report\EventSphere_Ch3.docx")

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)

def delete_table(table):
    t = table._element
    t.getparent().remove(t)

paragraphs = doc.paragraphs

# Map sections
sections = {}
for i, p in enumerate(paragraphs):
    txt = p.text.strip()
    if txt.startswith('3.2 Detailed Use Case'):
        sections['3.2_start'] = i
    elif txt.startswith('3.2.1'):
        sections['3.2.1_start'] = i
    elif txt.startswith('3.2.2'):
        sections['3.2.2_start'] = i
    elif txt.startswith('3.2.3'):
        sections['3.2.3_start'] = i
    elif txt.startswith('3.2.4'):
        sections['3.2.4_start'] = i
    elif txt.startswith('3.2.5'):
        sections['3.2.5_start'] = i
    elif txt.startswith('3.3 Functional'):
        sections['3.3_start'] = i

print(f"Found sections: {sections}")

# STEP 1: Replace all detailed use case descriptions (3.2.1-3.2.5) with a consolidated summary
# Delete everything from 3.2.1 to just before 3.3

to_delete = []
if '3.2.1_start' in sections and '3.3_start' in sections:
    for i in range(sections['3.2.1_start'], sections['3.3_start']):
        to_delete.append(i)

# Also delete the tables in that range
# Find all tables in the document and delete ones between use case sections
tables_to_delete = []
body = doc.element.body
all_elements = list(body)

# Track which tables are in the use case section
# Tables in docx are at the same level as paragraphs in the body
table_idx = 0
in_delete_zone = False
for elem in all_elements:
    if elem.tag == qn('w:p'):
        text = elem.text or ''
        # Also check runs
        for r in elem.findall('.//'+qn('w:t')):
            text += r.text or ''
        text = text.strip()
        if text.startswith('3.2.1'):
            in_delete_zone = True
        elif text.startswith('3.3 Functional'):
            in_delete_zone = False
    
    if elem.tag == qn('w:tbl') and in_delete_zone:
        tables_to_delete.append(elem)

# Delete tables in the use case zone
for tbl_elem in tables_to_delete:
    tbl_elem.getparent().remove(tbl_elem)
    print("Deleted a use case table")

# Now insert consolidated summary before deleting paragraphs
if '3.2.1_start' in sections:
    insert_ref = paragraphs[sections['3.2.1_start']]
    parent = insert_ref._element.getparent()
    insert_pos = list(parent).index(insert_ref._element)
    
    # Create consolidated content
    summary_parts = [
        ("3.2.1 Use Case Summary", "Heading3"),
        (
            "The following summarizes the five primary use cases of Event Sphere with their "
            "key actors, workflows and business rules:",
            "Normal"
        ),
        (
            "UC-1 User Management: The Admin manages all user accounts (create, update, delete) "
            "for Faculty and Student users. The system validates inputs, enforces email uniqueness, "
            "updates the database, and logs all operations. Only Admin users may execute user management "
            "operations and every account must have a unique ID with a verified institutional email.",
            "Normal"
        ),
        (
            "UC-2 Event Management: Faculty create and submit event proposals with full details. "
            "Submissions enter an Admin approval queue for review. The system manages validation, "
            "database storage, and automated notifications. Only approved events are visible to students, "
            "and rejected events must include mandatory Admin feedback. If modifications are needed, "
            "the event is returned to Faculty for resubmission.",
            "Normal"
        ),
        (
            "UC-3 Event Participation: Students browse approved events and register digitally through "
            "a one-tap process. On event day, students scan a unique QR code to confirm attendance. "
            "The system validates the scan (event ID match, registration check, duplicate prevention), "
            "records attendance, and marks certificate eligibility. Only registered students may scan "
            "the event QR code, and each QR code is unique per event and single-use per student.",
            "Normal"
        ),
        (
            "UC-4 Reports and Analytics: Admin and Faculty generate detailed reports on event activity, "
            "participation rates, attendance statistics, and departmental engagement. The system processes "
            "data into visual charts and exportable PDF/Excel formats. Faculty can only access analytics "
            "for events they created.",
            "Normal"
        ),
        (
            "UC-5 System Administration and Maintenance: The Admin maintains platform stability through "
            "system configuration updates, data backups, activity log reviews, user role management, "
            "and issue resolution. The system automates monitoring and scheduled backups. All system "
            "activities must be logged with a minimum 6-month retention period.",
            "Normal"
        ),
    ]
    
    # Insert in reverse order (since we insert at same position)
    for text, style_name in reversed(summary_parts):
        new_p = OxmlElement('w:p')
        new_ppr = OxmlElement('w:pPr')
        new_style = OxmlElement('w:pStyle')
        
        # Map style names
        style_val = style_name
        if style_name == 'Heading3':
            style_val = 'Heading3'
        
        new_style.set(qn('w:val'), style_val)
        new_ppr.append(new_style)
        new_p.append(new_ppr)
        new_run = OxmlElement('w:r')
        new_text = OxmlElement('w:t')
        new_text.set(qn('xml:space'), 'preserve')
        new_text.text = text
        new_run.append(new_text)
        new_p.append(new_run)
        parent.insert(insert_pos, new_p)

# Delete old use case paragraphs
deleted = 0
for i in sorted(to_delete, reverse=True):
    try:
        delete_paragraph(paragraphs[i])
        deleted += 1
    except:
        pass

print(f"Chapter 3: Deleted {deleted} paragraphs from use case section")

# STEP 2: Remove repetitive NFR intro sentences
# Re-read paragraphs after modifications
paragraphs2 = doc.paragraphs
nfr_deleted = 0
for p in paragraphs2:
    txt = p.text.strip()
    if txt.startswith("This section defines the") and "standards that Event Sphere must satisfy" in txt:
        delete_paragraph(p)
        nfr_deleted += 1

print(f"Chapter 3: Removed {nfr_deleted} repetitive NFR intro sentences")

doc.save(r"E:\FYP-main\report_extracted\report\EventSphere_Ch3.docx")
print("Chapter 3 saved successfully!")
